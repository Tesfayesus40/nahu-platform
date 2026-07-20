#!/usr/bin/env python3
"""Generate Amharic NAHU Transaction Protection Policy Word document."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

OUT = Path(__file__).resolve().parent / "Nahu_Transaction_Protection_Policy_v1_Amharic.docx"
AMHARIC_FONT = "Nyala"


def set_run(run, size=12, bold=False, color=None):
    run.font.name = AMHARIC_FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), AMHARIC_FONT)
    run._element.rPr.rFonts.set(qn("w:cs"), AMHARIC_FONT)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color


def heading(doc, text, level=1):
    p = doc.add_heading(level=level)
    run = p.add_run(text)
    set_run(run, size={1: 16, 2: 14, 3: 12}.get(level, 12), bold=True)
    return p


def para(doc, text, bold=False, size=12):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run(run, size=size, bold=bold)
    p.paragraph_format.space_after = Pt(6)
    return p


def quote(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.75)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    set_run(run, size=12, color=RGBColor(0x1A, 0x4D, 0x2E))
    run.italic = True
    return p


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    set_run(run, size=12)
    return p


def numbered(doc, text):
    p = doc.add_paragraph(style="List Number")
    run = p.add_run(text)
    set_run(run, size=12)
    return p


def flow(doc, lines):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run("\n".join(lines))
    set_run(run, size=11)
    return p


def table(doc, headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        set_run(run, size=11, bold=True)
    for r_i, row in enumerate(rows):
        for c_i, val in enumerate(row):
            cell = t.rows[r_i + 1].cells[c_i]
            cell.text = ""
            run = cell.paragraphs[0].add_run(val)
            set_run(run, size=11)
    doc.add_paragraph()
    return t


def build():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("ናሁ ፕላትፎርም")
    set_run(r, size=14, bold=True, color=RGBColor(0x1A, 0x4D, 0x2E))

    subt = doc.add_paragraph()
    subt.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subt.add_run("የግብይት ጥበቃ ፖሊሲ")
    set_run(r, size=20, bold=True)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = meta.add_run(
        "እምነት፣ ክወና እና ሕጋዊ ኃላፊነት\n"
        "ተፈፃሚነት፡ የገበሬ መተግበሪያ · የገዢ መተግበሪያ · ገበያ\n"
        "ሥሪት 1.0  |  ጁላይ 2026  |  ረቂቅ ለጅምር ግምገማ"
    )
    set_run(r, size=11)

    para(
        doc,
        "ይህ ሰነድ ገዢዎች እና ሻጮች ናሁን ከመተማመን በፊት የሚጠይቋቸውን ዋና ጥያቄዎች ይመልሳል፡ "
        "ለክፍያ ማን ኃላፊ ነው፣ ግብይቱን ምን ያረጋግጣል፣ እና ችግር ሲፈጠር ምን ይሆናል። "
        "መልሱ ዋናው ቴክኒካዊ አይደለም — እምነት፣ ክወና እና ሕጋዊ ኃላፊነት ነው።",
    )

    heading(doc, "፩. እያንዳንዱ ገዢ እና ሻጭ የሚጠይቋቸው ሶስት ጥያቄዎች", 1)

    heading(doc, "፩.፩ ለክፍያዬ ማን ኃላፊ ነው?", 2)
    para(doc, "ህዝባዊ መልስዎ ግልጽ መሆን አለበት፡")
    quote(
        doc,
        "ናሁ የግብይቱን ሂደት ለመቆጣጠር ኃላፊ ነው፣ ክፍያው ግን በተፈቀዱ የክፍያ ጣቢያዎች በኩል ይከናወናል። "
        "ናሁ ትዕዛዙን ይከታተላል፣ የክፍያ ሁኔታን ያረጋግጣል፣ እና ችግር ሲፈጠር ድጋፍ ይሰጣል።",
    )
    para(doc, "ውስጣዊ የኃላፊነት ክፍፍል፡", bold=True)
    table(
        doc,
        ["ኃላፊነት", "ባለቤት"],
        [
            ["የክፍያ ማካሄድ", "የክፍያ አቅራቢ (ቴሌብር፣ ሲቢኢ ብር፣ ባንክ)"],
            ["የትዕዛዝ አስተዳደር", "ናሁ ፕላትፎርም"],
            ["የክርክር አስተዳደር", "የናሁ ድጋፍ ቡድን"],
            ["የሻጭ (ገበሬ) ማረጋገጫ", "ናሁ"],
            ["የገዢ ማረጋገጫ", "ናሁ"],
            ["የግብይት መዝገቦች / ኦዲት", "ናሁ"],
        ],
    )
    para(
        doc,
        "ገንዘቡ በሦስተኛ ወገን የክፍያ አገልግሎት ቢንቀሳቀስም ተጠቃሚዎች ማንን ማነጋገር እንዳለባቸው ያውቃሉ። "
        "ናሁ ባንክ እንደሆነ አያቅርብም፤ የገበያ ግብይቱን የሂደት እና የድጋፍ ባለቤት ነው።",
    )

    heading(doc, "፩.፪ ክፍያዬን ምን ያረጋግጣል?", 2)
    para(
        doc,
        "ማረጋገጫው የክፍያ አቅራቢው ብቻ አይደለም። በእያንዳንዱ ደረጃ መዝገብ ያለው ሙሉ የግብይት ሂደት ነው።",
    )
    para(doc, "ዋና የግብይት ፍሰት (የቡና ገበያ አሁን)፡", bold=True)
    flow(
        doc,
        [
            "ገዢ ትዕዛዝ ያስቀምጣል",
            "        │",
            "ክፍያ ይቀርባል",
            "        │",
            "ክፍያ ይረጋገጣል",
            "        │",
            "ሻጭ (ገበሬ) ማሳወቂያ ይደርሰዋል",
            "        │",
            "ሻጭ ትዕዛዙን ይቀበላል / ያዘጋጃል",
            "        │",
            "ምርቱ ይላካል",
            "        │",
            "ገዢ መድረሱን ያረጋግጣል",
            "        │",
            "ግብይቱ ይጠናቀቃል",
        ],
    )
    para(doc, "በዚህ ሂደት ውስጥ ናሁ የሚያረጋግጠው፡", bold=True)
    bullet(doc, "እያንዳንዱ እርምጃ ይመዘገባል።")
    bullet(doc, "እያንዳንዱ የሁኔታ ለውጥ የጊዜ ማህተም ይኖረዋል።")
    bullet(doc, "እያንዳንዱ የክፍያ ማጣቀሻ ይቀመጣል።")
    bullet(doc, "እያንዳንዱ ክርክር ይከታተላል።")
    bullet(doc, "ድጋፍ ሙሉ የትዕዛዝ ታሪክን መገንባት ይችላል።")
    para(
        doc,
        "ይህ የኦዲት ዱካ ከግብይት ጥበቃው አካል ነው። ከደረጃ ያለው የትዕዛዝ ሁኔታ "
        "(ክፍያ በመንቀሳቀስ እስከ መድረስ ማረጋገጫ ድረስ በኤስክሮ ውስጥ እንደሚቆይ ተገንዝቦ) ጋር ሲጣመር "
        "ለሁለቱም ወገኖች ተግባራዊ ጥበቃ ይፈጥራል።",
    )

    heading(doc, "፩.፫ ግብይቱ ካልተሳካ ምን ይሆናል?", 2)
    para(doc, "ግልጽ የውድቀት ሁኔታዎች መገለፅ እና መታተም አለባቸው።", bold=True)

    heading(doc, "ሁኔታ ሀ — ገዢ ከፍሏል፣ ሻጭ አልላከም", 3)
    flow(
        doc,
        [
            "ተከፍሏል / በኤስክሮ ውስጥ",
            "        ↓",
            "የሻጭ ጊዜ አልፏል (በፖሊሲ ጊዜ ውስጥ አልተፈፀመም)",
            "        ↓",
            "ትዕዛዝ ተሰርዟል",
            "        ↓",
            "የገንዘብ መመለስ ሂደት ተጀምሯል",
        ],
    )

    heading(doc, "ሁኔታ ለ — ገዢ ምርቱ አልደረሰኝም ይላል", 3)
    flow(
        doc,
        [
            "ክርክር ተከፍቷል",
            "        ↓",
            "ማስረጃ ተሰብስቧል (ገዢ እና ሻጭ)",
            "        ↓",
            "የአስተዳዳሪ / ድጋፍ ምርመራ",
            "        ↓",
            "የተመዘገበ ውሳኔ",
            "        ↓",
            "ገንዘብ መመለስ  ወይም  ለገበሬ መለቀቅ",
        ],
    )

    heading(doc, "ሁኔታ ሐ — ገዢ የተሳሳተ መጠን ከፍሏል", 3)
    flow(
        doc,
        [
            "የክፍያ ማረጋገጫ",
            "        ↓",
            "አለመጣጣን ተገኝቷል",
            "        ↓",
            "ገዢ ተነግሮታል",
            "        ↓",
            "እርማት ተጠይቋል (እስኪፈታ ድረስ ለሻጭ አይለቀቅም)",
        ],
    )

    heading(doc, "ሁኔታ መ — የክፍያ ጣቢያ ስህተት", 3)
    flow(
        doc,
        [
            "ክፍያ አልተሳካም",
            "        ↓",
            "ትዕዛዝ በPENDING_PAYMENT ይቆያል",
            "        ↓",
            "ገዢ ክፍያን እንደገና ይሞክራል",
            "        ↓",
            "ክፍያ እስኪረጋገጥ ድረስ ለሻጭ ተጠባቂነት አይኖርም",
        ],
    )

    heading(doc, "፪. በናሁ ውስጥ ምን መከሰት አለበት", 1)
    para(
        doc,
        "ናሁ በይፋ በተቀመጠ የግብይት ጥበቃ ፖሊሲ መሥራት አለበት። እያንዳንዱ ትዕዛዝ ግልጽ ሁኔታዎችን "
        "ማለፍ አለበት ስለዚህ መተግበሪያዎች፣ ድጋፍ እና ተጠቃሚዎች አንድ ቋንቋ ይጋራሉ።",
    )

    heading(doc, "፪.፩ ዋና የትዕዛዝ ዑደት", 2)
    flow(
        doc,
        [
            "PENDING_PAYMENT — ክፍያ በመጠባበቅ ላይ",
            "        ↓",
            "PAYMENT_RECEIVED / PAID_ESCROW — ክፍያ ተቀብሏል / በኤስክሮ",
            "        ↓",
            "PAYMENT_VERIFIED — ክፍያ ተረጋግጧል",
            "        ↓",
            "SELLER_ACCEPTED — ሻጭ ተቀብሏል (አማራጭ)",
            "        ↓",
            "PREPARING — በመዘጋጀት ላይ",
            "        ↓",
            "SHIPPED — ተልኳል",
            "        ↓",
            "DELIVERED — ደርሷል",
            "        ↓",
            "BUYER_CONFIRMED — ገዢ አረጋግጧል",
            "        ↓",
            "FUNDS_RELEASED — ገንዘብ ተለቅቋል",
            "        ↓",
            "COMPLETED — ተጠናቅቋል",
        ],
    )
    para(
        doc,
        "ማሳሰቢያ — የአሁኑ የቡና ገበያ (ጁላይ 2026)፡ ቀላል ሁኔታዎችን ይጠቀማል፡ "
        "PENDING_PAYMENT → PAID_ESCROW → COMPLETED፣ ከCANCELLED እና DISPUTED ጋር። "
        "ከላይ ያለው ሙሉ ደረጃ የጭነት እና የፋይናንስ አገልግሎቶች ሲያድጉ የሚሄድ ግብ ነው።",
    )

    heading(doc, "፪.፪ የልዩ ሁኔታዎች ደረጃዎች", 2)
    flow(
        doc,
        [
            "PAYMENT_FAILED — ክፍያ አልተሳካም",
            "PAYMENT_EXPIRED — የክፍያ ጊዜ አልፏል",
            "CANCELLED — ተሰርዟል",
            "REFUND_PENDING — መመለስ በመጠባበቅ ላይ",
            "REFUNDED — ተመልሷል",
            "DISPUTED — በክርክር ላይ",
        ],
    )

    heading(doc, "፪.፫ ከአሁኑ ገበሬ እና ገዢ መተግበሪያዎች ጋር ማያያዝ", 2)
    table(
        doc,
        ["የተጠቃሚ እርምጃ (አሁን)", "በፖሊሲ ትርጉም"],
        [
            ["ገዢ ትዕዛዝ አስቀምጦ ይከፍላል (ቴሌብር / ሲቢኢ ብር)", "ክፍያ ቀርቧል → ተረጋግጧል → ገንዘብ በኤስክሮ ጽንሰ-ሐሳብ"],
            ["ገበሬ በትዕዛዝ ትር ያያል", "ሻጭ ተነግሮታል፤ ትዕዛዝ በጥበቃ ስር"],
            ["ገዢ መድረሱን ያረጋግጣል (አዎ / አይ)", "የገዢ ማረጋገጫ ወይም የክርክር ማስነሻ"],
            ["የመነሻ ምስክር ወረቀት ይሰጣል", "ለተሳካ ስምምነት ማስረጃ"],
            ["ክርክር ማንሳት (ሁለቱም ወገኖች)", "ጉዳዩ ወደ ናሁ ድጋፍ ለውሳኔ ይሄዳል"],
        ],
    )

    heading(doc, "፫. የመጨረሻ ውሳኔን ማን ይወስናል?", 1)
    para(doc, "በመጀመሪያ፡ የናሁ ድጋፍ ቡድን።", bold=True)
    para(doc, "ፕላትፎርሙ አስተዳዳሪዎች የሚከተሉትን የሚያደርጉበት የክርክር ዳሽቦርድ መኖር አለበት፡")
    bullet(doc, "ሙሉ የትዕዛዝ ጊዜ መስመር መመልከት።")
    bullet(doc, "የክፍያ ዝርዝሮችን እና ማጣቀሻዎችን መመርመር።")
    bullet(doc, "የተሰቀሉ ማስረጃዎችን ማየት።")
    bullet(doc, "የገዢ እና ሻጭ መልዕክቶችን ማንበብ።")
    bullet(doc, "የተመዘገበ ውሳኔ መመዝገብ።")
    para(doc, "እያንዳንዱ ውሳኔ ማከማቸት ያለበት፡", bold=True)
    bullet(doc, "የአስተዳዳሪ ማንነት")
    bullet(doc, "ቀን እና ሰዓት")
    bullet(doc, "ምክንያት / ማብራሪያ")
    bullet(doc, "የማስረጃ ማጣቀሻዎች")
    para(
        doc,
        "ይህ ኃላፊነትን ይፈጥራል እና እምነትን የሚጎዱ ያልተመዘገቡ ውሳኔዎችን ይቀንሳል።",
    )

    heading(doc, "፬. ተጠቃሚዎች ምን ማየት አለባቸው", 1)
    para(
        doc,
        "ገዢ ወይም ገበሬ «ክፍያዬን ምን ያረጋግጣል?» ሲሉ በመተግበሪያዎች እና በድረ-ገጽ ላይ "
        "የሚከተለው መልአክ ማሳየት አለበት፡",
    )
    quote(
        doc,
        "በናሁ ላይ እያንዳንዱ ግብይት ከክፍያ እስከ መድረስ ይቆጣጠራል። የክፍያ ሁኔታ፣ የትዕዛዝ ሁኔታ "
        "እና የመላኪያ ዝመናዎች በግብይቱ ሂደት ይመዘገባሉ። ችግር ሲፈጠር በፕላትፎርሙ በኩል ክርክር "
        "መክፈት ይችላሉ። የድጋፍ ቡድናችን ጉዳዩን ይመረምራል፣ ያሉ ማስረጃዎችን ይመለከታል፣ "
        "እና በናሁ የግብይት ፖሊሲ መሠረት ተገቢውን መፍትሄ ይወስናል።",
    )
    para(
        doc,
        "ይህ ተጨባጭ ተስፋ ያስቀምጣል እና ፕላትፎርሙ ሕጋዊ ማረጋገጥ የማይችለውን "
        "(ለምሳሌ ናሁ ፈቃድ ያለው ባንክ ነው ወይም ሁል ጊዜ ገንዘቡን ይይዛል) አያስተናግድም።",
    )

    heading(doc, "፭. የረጅም ጊዜ አርኪቴክቸር ምክር", 1)
    para(
        doc,
        "የግብይት ጥበቃ ከእርሻ፣ መላኪያ እና ፋይናንስ ጎን ለጎን የናሁ ዋና የንግድ ችሎታ መሆን አለበት።",
    )
    flow(
        doc,
        [
            "የግብይት ጥበቃ",
            "├── የክፍያ ማረጋገጫ",
            "├── የኤስክሮ ሁኔታ ክትትል",
            "├── የክርክር አስተዳደር",
            "├── የገንዘብ መመለስ አስተዳደር",
            "├── የማጭበርበር ፈልጎ ማውጣት",
            "├── የግብይት ኦዲት",
            "├── የገዢ ጥበቃ",
            "└── የሻጭ ጥበቃ",
        ],
    )
    para(
        doc,
        "ይህ ንግግሩን ከ«ገንዘቡን ማን ይይዛል?» ወደ «ናሁ ሁለቱንም ወገኖች እንዴት ይጠብቃል?» ይቀይራል። "
        "ይህ ጠንካራ የዋጋ ሀሳብ ነው እና የክፍያ ውህደቶች እና ክወናዎች ሲያድጉ ማደግ ይችላል።",
    )

    heading(doc, "፮. ከጅምር በፊት የማረጋገጫ ዝርዝር", 1)
    numbered(doc, "አጭር የግብይት ጥበቃ ማጠቃለያ በሁለቱም ገበሬ እና ገዢ መተግበሪያዎች ውስጥ ማተም (ቅንብሮች / እገዛ)።")
    numbered(doc, "ለክርክሮች የድጋፍ የጊዜ ሰሌዳ መወሰን (ለምሳሌ በ24 ሰዓት ውስጥ መቀበል፣ በX ቀን ውስጥ ውሳኔ)።")
    numbered(doc, "የሻጭ መሟላት ጊዜ እና የገዢ ማረጋገጫ መስኮት በጽሁፍ መግለፅ።")
    numbered(doc, "ከእያንዳንዱ የክፍያ አቅራቢ ጋር የገንዘብ መመለስ መንገድ ማረጋገጥ (ቴሌብር፣ ሲቢኢ ብር)።")
    numbered(doc, "እያንዳንዱ የድጋፍ ውሳኔ ከምክንያት እና ማስረጃ ጋር መመዝገቡን ማስገደድ።")
    numbered(doc, "ከህዝባዊ ጅምር በፊት ሕጋዊ / የሕግ አማካሪ ግምገማ ማድረግ።")
    numbered(doc, "ለአስተዳዳሪ የክርክር ዳሽቦርድ እና ሙሉ የሁኔታ ሞዴል የምርት እቅድ መቀጠል።")

    heading(doc, "፯. ተዛማጅ ስርዓቶች (የአሁኑ አተገባበር)", 1)
    table(
        doc,
        ["ስርዓት", "በግብይት ጥበቃ ውስጥ ያለው ሚና"],
        [
            ["nahu-platform ኤፒአይ", "ትዕዛዞች፣ የኤስክሮ ሁኔታ፣ ምስክር ወረቀቶች፣ የክርክር ምልክት፣ የክፍያ ማጣቀሻዎች"],
            ["የገበሬ መተግበሪያ", "ትዕዛዞችን መቀበል፣ ያልተከፈሉን መከልከል፣ ክርክር ማንሳት"],
            ["የገዢ መተግበሪያ", "መክፈል፣ መድረስ ማረጋገጥ፣ ያልተከፈለን መሰረዝ፣ ክርክር፣ ምስክር ወረቀት ማየት"],
            ["የክፍያ አቅራቢዎች", "ክፍያን ማስፈጸም፤ ናሁ ማጣቀሻን እና ውጤቱን ይከታተላል"],
            ["Staging / Production ዲቢ", "የኦዲት ዱካ ዋና መዝገብ"],
        ],
    )

    heading(doc, "፰. ማጠቃለያ", 1)
    para(
        doc,
        "ናሁን ከማስጀመር በፊት ድርጅቱ በ፩ ክፍል ያሉትን ሶስት ጥያቄዎች በአንድ ወጥ ድምፅ "
        "በድጋፍ፣ ምርት፣ ሕግ እና ኢንጂኔሪንግ መልሶ መመለስ መቻል አለበት። የግብይት ጥበቃ "
        "ናሁ እምነት የሚያገኝበት መንገድ ነው — ገንዘብን በመንቀሳቀሱ ብቻ ሳይሆን ሂደቱን፣ "
        "መዝገቦቹን እና ችግር ሲፈጠር መፍትሄውን በመያዙ።",
    )
    para(
        doc,
        "የሰነድ ባለቤት፡ ምርት / ክወና  ·  ገምጋሚዎች፡ ሕግ፣ ድጋፍ፣ ኢንጂኔሪንግ",
        size=10,
    )

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
