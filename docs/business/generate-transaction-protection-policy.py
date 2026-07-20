#!/usr/bin/env python3
"""Generate NAHU Transaction Protection Policy Word document."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

OUT = Path(__file__).resolve().parent / "Nahu_Transaction_Protection_Policy_v1.docx"


def set_run(run, size=11, bold=False, color=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color


def heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        set_run(run, size={1: 16, 2: 13, 3: 11}.get(level, 11), bold=True)
    return p


def para(doc, text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run(run, size=size, bold=bold)
    run.italic = italic
    p.paragraph_format.space_after = Pt(6)
    return p


def quote(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.75)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    set_run(run, size=11, color=RGBColor(0x1A, 0x4D, 0x2E))
    run.italic = True
    return p


def bullet(doc, text):
    p = doc.add_paragraph(text, style="List Bullet")
    for run in p.runs:
        set_run(run, size=11)
    return p


def numbered(doc, text):
    p = doc.add_paragraph(text, style="List Number")
    for run in p.runs:
        set_run(run, size=11)
    return p


def flow(doc, lines):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run("\n".join(lines))
    set_run(run, size=10)
    run.font.name = "Consolas"
    return p


def table(doc, headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            set_run(run, size=10, bold=True)
    for r_i, row in enumerate(rows):
        for c_i, val in enumerate(row):
            cell = t.rows[r_i + 1].cells[c_i]
            cell.text = val
            for run in cell.paragraphs[0].runs:
                set_run(run, size=10)
    doc.add_paragraph()
    return t


def build():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Cover / title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("NAHU PLATFORM")
    set_run(r, size=14, bold=True, color=RGBColor(0x1A, 0x4D, 0x2E))

    subt = doc.add_paragraph()
    subt.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subt.add_run("Transaction Protection Policy")
    set_run(r, size=20, bold=True)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = meta.add_run(
        "Trust, Operations & Legal Responsibility\n"
        "Applies to: Farmer App · Buyer App · Marketplace\n"
        "Version 1.0  |  July 2026  |  Draft for launch review"
    )
    set_run(r, size=10)

    para(
        doc,
        "This document answers the most important business questions buyers and sellers "
        "ask before trusting NAHU: who is responsible for payment, what guarantees a "
        "transaction, and what happens when something fails. The answer is not primarily "
        "technical—it is about trust, operations, and legal responsibility.",
    )

    # §1 Three questions
    heading(doc, "1. The Three Questions Every Buyer and Seller Will Ask", 1)

    heading(doc, "1.1 Who is responsible for my payment?", 2)
    para(
        doc,
        "Your public answer should be clear:",
    )
    quote(
        doc,
        "NAHU is responsible for managing the transaction process, while the payment is "
        "processed through approved payment channels. NAHU monitors the order, verifies "
        "transaction status, and provides support if any issue occurs.",
    )
    para(doc, "Internal responsibility split:", bold=True)
    table(
        doc,
        ["Responsibility", "Owner"],
        [
            ["Payment processing", "Payment provider (Telebirr, CBE Birr, bank)"],
            ["Order management", "NAHU Platform"],
            ["Dispute management", "NAHU Support Team"],
            ["Seller (farmer) verification", "NAHU"],
            ["Buyer verification", "NAHU"],
            ["Transaction records / audit trail", "NAHU"],
        ],
    )
    para(
        doc,
        "Users know who to contact even though a third-party payment service moves the money. "
        "NAHU does not present itself as the bank; NAHU presents itself as the process owner "
        "and support owner for the marketplace transaction.",
    )

    heading(doc, "1.2 What guarantees my payment?", 2)
    para(
        doc,
        "The guarantee is not only the payment provider. It is the complete transaction "
        "workflow, with records at every step.",
    )
    para(doc, "Core transaction flow (marketplace — coffee apps today):", bold=True)
    flow(
        doc,
        [
            "Buyer places order",
            "        │",
            "Payment submitted",
            "        │",
            "Payment verified",
            "        │",
            "Seller (farmer) notified",
            "        │",
            "Seller accepts / prepares order",
            "        │",
            "Product delivered",
            "        │",
            "Buyer confirms delivery",
            "        │",
            "Transaction completed",
        ],
    )
    para(doc, "Throughout this process NAHU ensures that:", bold=True)
    bullet(doc, "Every action is recorded.")
    bullet(doc, "Every status change is timestamped.")
    bullet(doc, "Every payment reference is stored.")
    bullet(doc, "Every dispute is tracked.")
    bullet(doc, "Support can reconstruct the full timeline of the order.")
    para(
        doc,
        "That audit trail is part of the guarantee. Combined with staged order statuses "
        "(payment held conceptually in escrow until delivery confirmation), it forms the "
        "practical protection layer for both parties.",
    )

    heading(doc, "1.3 What if the transaction fails?", 2)
    para(doc, "Clear failure scenarios must be defined and published.", bold=True)

    heading(doc, "Scenario A — Buyer paid, seller never ships", 3)
    flow(
        doc,
        [
            "PAID / IN ESCROW",
            "        ↓",
            "Seller timeout (no fulfilment within policy window)",
            "        ↓",
            "Order cancelled",
            "        ↓",
            "Refund process initiated",
        ],
    )

    heading(doc, "Scenario B — Buyer claims product not received", 3)
    flow(
        doc,
        [
            "Dispute opened",
            "        ↓",
            "Evidence collected (buyer & seller)",
            "        ↓",
            "Admin / Support investigation",
            "        ↓",
            "Documented decision",
            "        ↓",
            "Refund  OR  Release payment to farmer",
        ],
    )

    heading(doc, "Scenario C — Buyer paid the wrong amount", 3)
    flow(
        doc,
        [
            "Payment verification",
            "        ↓",
            "Mismatch detected",
            "        ↓",
            "Buyer notified",
            "        ↓",
            "Correction requested (no seller release until resolved)",
        ],
    )

    heading(doc, "Scenario D — Payment gateway error", 3)
    flow(
        doc,
        [
            "Payment failed",
            "        ↓",
            "Order remains PENDING_PAYMENT",
            "        ↓",
            "Buyer retries payment",
            "        ↓",
            "No seller fulfilment expectation until payment is confirmed",
        ],
    )

    # §2 Inside NAHU
    heading(doc, "2. What Should Happen Inside NAHU", 1)
    para(
        doc,
        "NAHU should operate under a formal Transaction Protection Policy. Every order "
        "must move through clear statuses so apps, support, and users share one language.",
    )

    heading(doc, "2.1 Primary order lifecycle", 2)
    flow(
        doc,
        [
            "PENDING_PAYMENT",
            "        ↓",
            "PAYMENT_RECEIVED / PAID_ESCROW",
            "        ↓",
            "PAYMENT_VERIFIED",
            "        ↓",
            "SELLER_ACCEPTED (optional gate)",
            "        ↓",
            "PREPARING / IN_FULFILMENT",
            "        ↓",
            "SHIPPED / IN_TRANSIT",
            "        ↓",
            "DELIVERED",
            "        ↓",
            "BUYER_CONFIRMED",
            "        ↓",
            "FUNDS_RELEASED",
            "        ↓",
            "COMPLETED",
        ],
    )
    para(
        doc,
        "Note — Current coffee marketplace (July 2026): the live apps use a simplified "
        "subset: PENDING_PAYMENT → PAID_ESCROW → COMPLETED, plus CANCELLED and DISPUTED. "
        "The fuller ladder above is the target model as logistics and finance mature.",
    )

    heading(doc, "2.2 Exception states", 2)
    flow(
        doc,
        [
            "PAYMENT_FAILED",
            "PAYMENT_EXPIRED",
            "CANCELLED",
            "REFUND_PENDING",
            "REFUNDED",
            "DISPUTED",
        ],
    )

    heading(doc, "2.3 Mapping to today’s farmer & buyer apps", 2)
    table(
        doc,
        ["User action (today)", "Policy meaning"],
        [
            ["Buyer places order & pays (Telebirr / CBE Birr)", "Payment submitted → verified → funds conceptually in escrow"],
            ["Farmer sees order on Orders tab", "Seller notified; order under protection"],
            ["Buyer confirms delivery (Yes / No)", "Buyer confirmation or dispute trigger"],
            ["Origin certificate issued", "Proof for completed successful trade"],
            ["Raise dispute (either party)", "Case moves to NAHU Support for decision"],
        ],
    )

    # §3 Who decides
    heading(doc, "3. Who Makes the Final Decision?", 1)
    para(doc, "Initially: NAHU Support Team.", bold=True)
    para(
        doc,
        "The platform should provide an Admin Dispute Dashboard where administrators can:",
    )
    bullet(doc, "View the complete order timeline.")
    bullet(doc, "Review payment details and payment references.")
    bullet(doc, "See any uploaded evidence.")
    bullet(doc, "Read buyer and seller communications.")
    bullet(doc, "Record a documented decision.")
    para(doc, "Every decision must store:", bold=True)
    bullet(doc, "Administrator identity")
    bullet(doc, "Date and time")
    bullet(doc, "Reason / rationale")
    bullet(doc, "Supporting evidence references")
    para(
        doc,
        "This creates accountability and reduces informal, untracked decisions that undermine trust.",
    )

    # §4 What users see
    heading(doc, "4. What Users Should See", 1)
    para(
        doc,
        "When a buyer or farmer asks “What guarantees my payment?”, the apps and "
        "website should display language similar to:",
    )
    quote(
        doc,
        "Every transaction on NAHU is monitored from payment to delivery. Payment status, "
        "order status, and delivery updates are tracked throughout the transaction. If a "
        "problem occurs, you can open a dispute through the platform. Our support team "
        "will review the case, examine the available evidence, and determine the "
        "appropriate resolution according to NAHU’s transaction policy.",
    )
    para(
        doc,
        "This sets realistic expectations without promising something the platform "
        "cannot legally guarantee (for example, that NAHU itself is a licensed bank "
        "or holds settlement funds in all cases).",
    )

    # §5 Long-term
    heading(doc, "5. Long-Term Architecture Recommendation", 1)
    para(
        doc,
        "Transaction Protection should become a core business capability of NAHU, "
        "alongside Farms, Delivery, and Finance.",
    )
    flow(
        doc,
        [
            "Transaction Protection",
            "├── Payment Verification",
            "├── Escrow Status Tracking",
            "├── Dispute Management",
            "├── Refund Management",
            "├── Fraud Detection",
            "├── Transaction Audit",
            "├── Buyer Protection",
            "└── Seller Protection",
        ],
    )
    para(
        doc,
        "This shifts the conversation from “Who holds the money?” to “How does NAHU "
        "protect both parties?” That is a stronger value proposition and one that can "
        "evolve as payment integrations and operations mature.",
    )

    # §6 Launch checklist
    heading(doc, "6. Pre-Launch Checklist", 1)
    numbered(doc, "Publish a short Transaction Protection summary in both farmer and buyer apps (Settings / Help).")
    numbered(doc, "Document Support SLAs for disputes (e.g. acknowledge within 24h, decide within X days).")
    numbered(doc, "Define seller fulfilment timeout and buyer confirmation window in writing.")
    numbered(doc, "Confirm refund path with each payment provider (Telebirr, CBE Birr).")
    numbered(doc, "Require that every Support decision is logged with reason and evidence.")
    numbered(doc, "Align legal/counsel review of wording before public launch.")
    numbered(doc, "Keep product roadmap items for Admin Dispute Dashboard and fuller status model.")

    # §7 Related systems
    heading(doc, "7. Related Systems (Current Implementation)", 1)
    table(
        doc,
        ["System", "Role in Transaction Protection"],
        [
            ["nahu-platform API", "Orders, escrow status, certificates, dispute flag, payment references"],
            ["Farmer app", "Receive orders, decline unpaid orders, raise dispute"],
            ["Buyer app", "Pay, confirm delivery, cancel unpaid, raise dispute, view certificate"],
            ["Payment providers", "Execute payment; NAHU tracks reference and outcome"],
            ["Staging / Production DB", "System of record for audit trail"],
        ],
    )

    # Closing
    heading(doc, "8. Closing Statement", 1)
    para(
        doc,
        "Before launching NAHU, the organization must be able to answer the three "
        "questions in Section 1 in one consistent voice across Support, Product, Legal, "
        "and Engineering. Transaction Protection is how NAHU earns trust—not only by "
        "moving money, but by owning the process, the records, and the resolution when "
        "something goes wrong.",
    )
    para(
        doc,
        "Document owner: Product / Operations  ·  Reviewers: Legal, Support, Engineering",
        italic=True,
        size=10,
    )

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
